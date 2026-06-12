"""Email attachment routes (list/download attachments, attachment-as-doc).

register_attachment_routes(router) registers the /attachments and /attachment*
endpoints on the given router. Extracted verbatim from setup_email_routes() in
email_routes.py (Phase 2.2 / ADR-039); these routes use the module-level
_imap() context manager + email_helpers attachment functions, not the bound
pool/sync locals, so the registrar needs only the router.
"""

import email as email_mod
import json
import logging
import uuid
from datetime import datetime

from fastapi import Depends, Query, Request
from fastapi.responses import FileResponse

from routes.email_helpers import (
    require_owner,
    _imap,
    _q,
    _decode_header,
    _extract_attachment_text,
    _list_attachments_from_msg,
    _extract_attachment_to_disk,
    _extract_html,
    _extract_text,
    ATTACHMENTS_DIR,
)
from routes.email_route_helpers import _imap_uid_fetch

logger = logging.getLogger(__name__)


def register_attachment_routes(router):
    """Register the attachment endpoints on `router`."""
    @router.get("/attachments/{uid}")
    async def list_attachments(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """List attachments for an email."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder), readonly=True)
                status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
            if status != "OK":
                return {"attachments": [], "error": "Email not found"}
            raw = msg_data[0][1]
            msg = email_mod.message_from_bytes(raw)
            attachments = _list_attachments_from_msg(msg)
            return {"attachments": attachments, "uid": uid}
        except Exception as e:
            logger.error(f"Failed to list attachments for {uid}: {e}")
            return {"attachments": [], "error": "Mail operation failed"}

    @router.get("/attachment/{uid}/{index}")
    async def download_attachment(uid: str, index: int, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Download a specific attachment by email UID and attachment index. Saves to local disk and returns the file."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder), readonly=True)
                status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
            if status != "OK":
                return {"error": "Email not found"}
            raw = msg_data[0][1]
            msg = email_mod.message_from_bytes(raw)

            # Extract to a per-email folder
            target_dir = ATTACHMENTS_DIR / f"{folder}_{uid}"
            filepath = _extract_attachment_to_disk(msg, index, target_dir)
            if not filepath:
                return {"error": f"Attachment index {index} not found"}

            return FileResponse(
                path=str(filepath),
                filename=filepath.name,
                media_type="application/octet-stream",
            )
        except Exception as e:
            logger.error(f"Failed to download attachment {uid}/{index}: {e}")
            return {"error": "Mail operation failed"}

    @router.post("/attachment-as-doc/{uid}/{index}")
    async def attachment_as_doc(uid: str, index: int, request: Request, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Extract an email attachment and open it in the document editor.

        Supported extensions:
          - .pdf   → rendered as PDF Document (existing flow)
          - .docx  → text extracted to markdown Document
          - .txt / .md → loaded directly as a markdown Document

        Returns {doc_id} so the frontend can open it as a tab in the doc panel.
        Other types are rejected — caller should fall back to download.
        """
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder), readonly=True)
                status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
            if status != "OK":
                return {"error": "Email not found"}
            raw = msg_data[0][1]
            msg = email_mod.message_from_bytes(raw)

            target_dir = ATTACHMENTS_DIR / f"{folder}_{uid}"
            filepath = _extract_attachment_to_disk(msg, index, target_dir)
            if not filepath:
                return {"error": f"Attachment index {index} not found"}

            from pathlib import Path as _Path
            base = _Path(filepath).name
            if base.startswith("."):
                return {"error": "Invalid filename", "filename": base}
            ext = _Path(base).suffix.lower()

            import os as _os
            title = _os.path.splitext(filepath.name)[0]

            # Capture the source email's identity so the doc can later be used
            # to thread a signed-reply back to the original sender.
            src_message_id = (msg.get("Message-ID") or "").strip()
            def _tag_doc_with_source(doc_id_to_tag: str):
                if not doc_id_to_tag:
                    return
                try:
                    from src.database import SessionLocal as _SL, Document as _Doc
                    _db = _SL()
                    try:
                        d = _db.query(_Doc).filter(_Doc.id == doc_id_to_tag).first()
                        if d:
                            d.source_email_uid = str(uid)
                            d.source_email_folder = folder
                            d.source_email_account_id = account_id or ""
                            d.source_email_message_id = src_message_id
                            _db.commit()
                    finally:
                        _db.close()
                except Exception as _e:
                    logger.warning(f"tag doc source-email failed: {_e}")

            # Extracted docs MUST belong to a session the caller owns — a
            # session-less ("orphan") doc is rejected by get_document's owner
            # check (404), so the frontend's loadDocument() throws and nothing
            # opens (the "open in document didn't open" bug). Attach it to the
            # user's most-recent session so it's fetchable + ownable.
            from src.auth_helpers import get_current_user as _gcu
            _doc_user = _gcu(request)
            def _resolve_doc_session():
                try:
                    from src.database import SessionLocal as _SL, Session as _Sess
                    _db = _SL()
                    try:
                        _q2 = _db.query(_Sess)
                        if _doc_user:
                            _q2 = _q2.filter(_Sess.owner == _doc_user)
                        s = _q2.order_by(_Sess.updated_at.desc()).first()
                        return s.id if s else None
                    finally:
                        _db.close()
                except Exception as _e:
                    logger.warning(f"resolve doc session failed: {_e}")
                    return None
            doc_session_id = _resolve_doc_session()

            # ── PDF path (existing) ────────────────────────────────────
            if ext == ".pdf":
                import shutil as _shutil
                from src.constants import UPLOAD_DIR
                from src.pdf_forms import has_form_fields, extract_fields
                from src.pdf_form_doc import (
                    save_field_sidecar,
                    create_form_markdown_document,
                    create_plain_pdf_document,
                )

                upload_id = f"{uuid.uuid4().hex}.pdf"
                today = datetime.utcnow().strftime("%Y/%m/%d")
                dated_dir = _os.path.join(UPLOAD_DIR, today)
                _os.makedirs(dated_dir, exist_ok=True)
                dest_path = _os.path.join(dated_dir, upload_id)
                _shutil.copyfile(str(filepath), dest_path)

                is_form = False
                try:
                    is_form = has_form_fields(dest_path)
                except Exception as e:
                    logger.warning(f"has_form_fields failed for attachment PDF: {e}")

                if is_form:
                    fields = extract_fields(dest_path)
                    save_field_sidecar(dest_path, fields)
                    doc_id = create_form_markdown_document(
                        session_id=doc_session_id,
                        fields=fields,
                        upload_id=upload_id,
                        title=title,
                        intro_text=None,
                    )
                else:
                    doc_id = create_plain_pdf_document(
                        session_id=doc_session_id,
                        upload_id=upload_id,
                        title=title,
                    )

                if not doc_id:
                    return {"error": "Failed to create document"}
                _tag_doc_with_source(doc_id)
                return {"doc_id": doc_id, "filename": filepath.name}

            # ── DOCX path: extract text → markdown document ───────────
            if ext == ".docx":
                try:
                    from docx import Document as _Docx
                except ImportError:
                    return {"error": "python-docx not installed", "filename": base}
                try:
                    d = _Docx(str(filepath))
                except Exception as e:
                    return {"error": f"Failed to read docx: {e}", "filename": base}
                # Convert paragraphs to markdown — preserve heading styles as #/##/###,
                # bullet lists as `- `, numbered lists as `1.`, and keep tables as
                # simple pipe-delimited rows.
                lines: list[str] = []
                for p in d.paragraphs:
                    text = p.text or ""
                    style = (p.style.name if p.style else "") or ""
                    if not text.strip():
                        lines.append("")
                        continue
                    if style.startswith("Heading 1"): lines.append(f"# {text}")
                    elif style.startswith("Heading 2"): lines.append(f"## {text}")
                    elif style.startswith("Heading 3"): lines.append(f"### {text}")
                    elif style.startswith("Heading "): lines.append(f"#### {text}")
                    elif style.startswith("List Bullet"): lines.append(f"- {text}")
                    elif style.startswith("List Number"): lines.append(f"1. {text}")
                    else: lines.append(text)
                for tbl in d.tables:
                    lines.append("")
                    for ri, row in enumerate(tbl.rows):
                        cells = [(c.text or "").replace("|", "\\|").replace("\n", " ").strip() for c in row.cells]
                        lines.append("| " + " | ".join(cells) + " |")
                        if ri == 0:
                            lines.append("|" + "|".join(["---"] * len(cells)) + "|")
                    lines.append("")
                content = "\n".join(lines).strip() or f"_(empty {base})_"

                from src.database import SessionLocal as _SL, Document as _Doc, DocumentVersion as _DV
                doc_id = str(uuid.uuid4())
                ver_id = str(uuid.uuid4())
                _db = _SL()
                try:
                    _db.query(_Doc).filter(_Doc.is_active == True).update({"is_active": False})
                    _db.add(_Doc(
                        id=doc_id, session_id=doc_session_id, title=title,
                        language="markdown", current_content=content,
                        version_count=1, is_active=True,
                    ))
                    _db.add(_DV(
                        id=ver_id, document_id=doc_id, version_number=1,
                        content=content, summary="Imported from DOCX", source="upload",
                    ))
                    _db.commit()
                finally:
                    _db.close()
                _tag_doc_with_source(doc_id)
                return {"doc_id": doc_id, "filename": filepath.name}

            # ── Plain text / markdown ────────────────────────────────
            if ext in (".txt", ".md", ".markdown"):
                try:
                    content = filepath.read_text(encoding="utf-8", errors="replace")
                except Exception as e:
                    return {"error": f"Failed to read text file: {e}", "filename": base}
                from src.database import SessionLocal as _SL, Document as _Doc, DocumentVersion as _DV
                doc_id = str(uuid.uuid4())
                ver_id = str(uuid.uuid4())
                _db = _SL()
                try:
                    _db.query(_Doc).filter(_Doc.is_active == True).update({"is_active": False})
                    _db.add(_Doc(
                        id=doc_id, session_id=doc_session_id, title=title,
                        language="markdown", current_content=content,
                        version_count=1, is_active=True,
                    ))
                    _db.add(_DV(
                        id=ver_id, document_id=doc_id, version_number=1,
                        content=content, summary="Imported from email attachment", source="upload",
                    ))
                    _db.commit()
                finally:
                    _db.close()
                _tag_doc_with_source(doc_id)
                return {"doc_id": doc_id, "filename": filepath.name}

            return {"error": f"Unsupported attachment type: {ext}", "filename": base}
        except Exception as e:
            logger.error(f"attachment-as-doc {uid}/{index} failed: {e}")
            return {"error": "Mail operation failed"}

    @router.post("/attachment-path/{uid}/{index}")
    async def get_attachment_path(uid: str, index: int, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
        """Extract attachment to local disk and return the path (for AI to read via read_file)."""
        try:
            with _imap(account_id, owner=owner) as conn:
                conn.select(_q(folder), readonly=True)
                status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
            if status != "OK":
                return {"error": "Email not found"}
            raw = msg_data[0][1]
            msg = email_mod.message_from_bytes(raw)

            target_dir = ATTACHMENTS_DIR / f"{folder}_{uid}"
            filepath = _extract_attachment_to_disk(msg, index, target_dir)
            if not filepath:
                return {"error": f"Attachment index {index} not found"}

            return {"path": str(filepath), "filename": filepath.name, "size": filepath.stat().st_size}
        except Exception as e:
            logger.error(f"Failed to get attachment path {uid}/{index}: {e}")
            return {"error": "Mail operation failed"}
